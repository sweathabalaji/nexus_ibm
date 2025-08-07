import os
import json
from typing import List, Dict, Any, Optional
from pydantic import BaseModel, Field
from langchain_groq import ChatGroq
from langchain_tavily import TavilySearch
from langchain_core.messages import HumanMessage
from dotenv import load_dotenv
import PyPDF2
from docx import Document
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse
import tempfile
import uvicorn
import re
import sys
from pathlib import Path
import requests

# Load environment variables from .env file
load_dotenv()

# Initialize FastAPI app
app = FastAPI(title="Resume Analyzer API", version="1.0.0")

# --- 1. Pydantic Schemas ---

class CompanyDetail(BaseModel):
    """Structured details for a single company."""
    company_name: str = Field(description="The name of the company.")
    legitimacy_status: str = Field(description="Legitimacy assessment (e.g., 'Legitimate', 'Suspicious', 'Unknown').")
    payment_status: str = Field(description="Payment structure (e.g., 'Paid', 'Unpaid', 'Student Paid Fee', 'Unknown').")
    key_findings: str = Field(description="Brief summary of evidence and findings.")
    reputation_score: int = Field(description="Score from 1-10 indicating company reputation.")

class JobMatch(BaseModel):
    """Resume-to-job-description match analysis."""
    match_score: float = Field(description="Score out of 100 for resume-job match.")
    summary: str = Field(description="Summary explaining the score and key points.")

class UnifiedReport(BaseModel):
    """Complete analysis report."""
    job_match_analysis: JobMatch
    company_verification: List[CompanyDetail]

class AnalysisResponse(BaseModel):
    """FastAPI response model."""
    company_verification: List[CompanyDetail]
    job_match_analysis: JobMatch

class TextAnalysisRequest(BaseModel):
    """Request model for text analysis."""
    resume_text: str = Field(..., description="Resume text content")
    job_description_text: str = Field(..., description="Job description text content")

# Resume Analysis Schemas
class ATSScoreBreakdown(BaseModel):
    """ATS score breakdown."""
    ats_score: int = Field(description="Overall ATS score (0-100)")
    formatting_score: int = Field(description="Formatting score (0-25)")
    keyword_density: int = Field(description="Keyword density score (0-25)")
    section_organization: int = Field(description="Section organization score (0-25)")
    readability: int = Field(description="Readability score (0-25)")

class KeywordAnalysis(BaseModel):
    """Keyword analysis results."""
    technical_skills: List[str] = Field(description="Technical skills found")
    soft_skills: List[str] = Field(description="Soft skills found")
    industry_keywords: List[str] = Field(description="Industry-specific keywords")
    missing_keywords: List[str] = Field(description="Missing important keywords")
    keyword_density_score: int = Field(description="Keyword density score (0-100)")

class ActionVerbAnalysis(BaseModel):
    """Action verb analysis results."""
    strong_verbs: List[str] = Field(description="Strong action verbs found")
    weak_verbs: List[str] = Field(description="Weak action verbs found")
    verb_strength_score: int = Field(description="Verb strength score (0-100)")
    verb_variety_score: int = Field(description="Verb variety score (0-100)")

class StrengthsWeaknesses(BaseModel):
    """Strengths and weaknesses analysis."""
    strengths: List[str] = Field(description="Resume strengths")
    weaknesses: List[str] = Field(description="Resume weaknesses")
    overall_impression: str = Field(description="Overall impression")
    competitiveness_score: int = Field(description="Competitiveness score (0-100)")

class ImprovementRecommendation(BaseModel):
    """Individual improvement recommendation."""
    category: str = Field(description="Category of improvement")
    issue: str = Field(description="Issue identified")
    recommendation: str = Field(description="Specific recommendation")
    priority: str = Field(description="Priority level (High/Medium/Low)")
    example: Optional[str] = Field(description="Example of improvement")

class ResumeAnalysisResult(BaseModel):
    """Complete resume analysis result."""
    file_path: str = Field(description="Path to analyzed file")
    analysis_timestamp: str = Field(description="Analysis timestamp")
    ats_score: int = Field(description="Overall ATS score")
    ats_details: ATSScoreBreakdown = Field(description="Detailed ATS breakdown")
    keyword_analysis: KeywordAnalysis = Field(description="Keyword analysis")
    action_verb_analysis: ActionVerbAnalysis = Field(description="Action verb analysis")
    strengths_weaknesses: StrengthsWeaknesses = Field(description="Strengths and weaknesses")
    priority_improvements: List[ImprovementRecommendation] = Field(description="Priority improvements")
    overall_grade: str = Field(description="Overall grade")

# Company Legitimacy Schemas
class CompanyAssessment(BaseModel):
    """Company assessment details."""
    one_line_summary: str = Field(description="Brief company description")
    legitimacy_status: str = Field(description="Legitimacy status")
    employer_review: str = Field(description="Employee reviews summary")
    growth: str = Field(description="Company growth status")
    revenue: str = Field(description="Revenue information")
    scale: str = Field(description="Company scale and size")
    job_role_assessment: Optional[str] = Field(description="Job role specific assessment")
    red_flags: str = Field(description="Warning signs")
    positive_indicators: str = Field(description="Positive indicators")

class CompanyLegitimacyResult(BaseModel):
    """Company legitimacy check result."""
    company_name: str = Field(description="Company name")
    job_role: Optional[str] = Field(description="Job role")
    assessment: CompanyAssessment = Field(description="Assessment details")
    status: str = Field(description="Check status")

class CompanyLegitimacyRequest(BaseModel):
    """Request for company legitimacy check."""
    company_name: str = Field(..., description="Company name to check")
    job_role: Optional[str] = Field(None, description="Optional job role")

class ResumeAnalysisRequest(BaseModel):
    """Request for resume analysis."""
    resume_text: str = Field(..., description="Resume text content")


# --- 2. File Processing Utilities ---

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF file."""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        doc = Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

def process_resume_file(file_path: str) -> str:
    """Process a resume file (PDF or DOCX) and extract text."""
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return extract_text_from_docx(file_path)
    else:
        return f"Unsupported file type: {file_path}"

def extract_text_from_file_bytes(file_content: bytes, filename: str) -> str:
    """Extract text from uploaded file bytes (PDF or DOCX)."""
    try:
        if filename.lower().endswith('.pdf'):
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(file_content)
                tmp_file.flush()
                
                with open(tmp_file.name, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                
                os.unlink(tmp_file.name)
                return text.strip()
                
        elif filename.lower().endswith('.docx'):
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(file_content)
                tmp_file.flush()
                
                doc = Document(tmp_file.name)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                os.unlink(tmp_file.name)
                return text.strip()
        else:
            return f"Unsupported file type: {filename}"
    except Exception as e:
        return f"Error extracting text from {filename}: {str(e)}"


# --- 3. AI Service Classes ---

class MoonshotAI:
    """Custom LLM wrapper for MoonshotAI API (via Groq)"""
    
    def __init__(self, api_key: str = None):
        self.api_key = api_key or os.getenv("MOONSHOT_API_KEY") or os.getenv("GROQ_API_KEY")
        self.model_name = "moonshotai/kimi-k2-instruct"
        self.base_url = "https://api.groq.com/openai/v1"
        
        if not self.api_key:
            raise ValueError("API key is required")
    
    def _call(self, prompt: str) -> str:
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": self.model_name,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3,
            "max_tokens": 4000
        }
        
        try:
            response = requests.post(
                f"{self.base_url}/chat/completions",
                headers=headers,
                json=data,
                timeout=60
            )
            response.raise_for_status()
            result = response.json()
            return result["choices"][0]["message"]["content"]
        except Exception as e:
            return f"Error calling MoonshotAI: {str(e)}"

class CompanyLegitimacyChecker:
    """Company legitimacy verification service."""
    
    def __init__(self):
        """Initialize the company legitimacy checker."""
        self.llm = ChatGroq(
            model="llama3-8b-8192", 
            groq_api_key=os.getenv("GROQ_API_KEY"),
            temperature=0.1
        )
        
        self.search_tool = TavilySearch(tavily_api_key=os.getenv("TAVILY_API_KEY"))
    
    def check_company_legitimacy(self, company_name: str, job_role: str = None) -> Dict[str, Any]:
        """Check if a company is legitimate and return comprehensive assessment."""
        try:
            # Search for company information
            search_query = f"{company_name} company legitimacy reviews scam check reputation employer review growth revenue scale size"
            search_results = self.search_tool.invoke(search_query)
            
            # Search for job role specific information if provided
            job_role_info = ""
            if job_role:
                job_search_query = f"{company_name} {job_role} employee reviews salary benefits work culture"
                job_search_results = self.search_tool.invoke(job_search_query)
                job_role_info = f"\n\nJob Role Specific Information for '{job_role}':\n{job_search_results}"
            
            verification_prompt = f"""
            Based on these search results about '{company_name}', provide a comprehensive assessment in JSON format:
            
            Search Results:
            {search_results}{job_role_info}
            
            Provide the following information in JSON format:
            {{
                "one_line_summary": "Brief one-line description of what the company does",
                "legitimacy_status": "Legitimate/Suspicious/Scam/Insufficient Data",
                "employer_review": "Overall rating and employee feedback summary",
                "growth": "Company growth status and trajectory",
                "revenue": "Revenue information or financial status",
                "scale": "Company size, number of employees, market presence",
                "job_role_assessment": "Specific assessment for the {job_role if job_role else 'requested'} role (if applicable)",
                "red_flags": "Any warning signs or concerns",
                "positive_indicators": "Good signs about the company"
            }}
            
            Ensure the response is valid JSON format. If information is not available, use "Not available" or "Insufficient data".
            """
            
            response = self.llm.invoke(verification_prompt)
            assessment_text = response.content if hasattr(response, 'content') else str(response)
            
            # Try to parse JSON from the response
            try:
                json_match = re.search(r'\{[\s\S]*\}', assessment_text)
                if json_match:
                    assessment_json = json.loads(json_match.group(0))
                else:
                    # If no JSON found, create structured response
                    assessment_json = {
                        "one_line_summary": "Unable to extract structured data",
                        "legitimacy_status": "Insufficient Data",
                        "employer_review": assessment_text[:200] + "...",
                        "growth": "Not available",
                        "revenue": "Not available", 
                        "scale": "Not available",
                        "job_role_assessment": "Not available" if not job_role else f"No specific data for {job_role}",
                        "red_flags": "Unable to determine",
                        "positive_indicators": "Unable to determine"
                    }
            except json.JSONDecodeError:
                assessment_json = {
                    "one_line_summary": "Data parsing error",
                    "legitimacy_status": "Insufficient Data",
                    "employer_review": assessment_text[:200] + "...",
                    "growth": "Not available",
                    "revenue": "Not available",
                    "scale": "Not available", 
                    "job_role_assessment": "Not available" if not job_role else f"No specific data for {job_role}",
                    "red_flags": "Unable to determine",
                    "positive_indicators": "Unable to determine"
                }
            
            return {
                "company_name": company_name,
                "job_role": job_role,
                "assessment": assessment_json,
                "status": "success"
            }
            
        except Exception as e:
            return {
                "company_name": company_name,
                "job_role": job_role,
                "assessment": f"Error checking legitimacy: {str(e)}",
                "status": "error"
            }

class DetailedResumeAnalyzer:
    """Detailed resume analyzer with ATS scoring (from RA.py)"""
    
    def __init__(self, api_key: str = None):
        self.llm = MoonshotAI(api_key=api_key)
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text content from PDF file"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            return f"Error extracting PDF: {str(e)}"
    
    def calculate_ats_score(self, resume_text: str) -> Dict[str, Any]:
        """Calculate ATS score using AI analysis"""
        prompt = f"""
        Analyze this resume text for ATS (Applicant Tracking System) compatibility and provide a comprehensive score.
        
        Resume Text:
        {resume_text}
        
        Please analyze and return a JSON response with the following structure:
        {{
            "ats_score": <number between 0-100>,
            "ats_factors": {{
                "formatting_score": <0-25>,
                "keyword_density": <0-25>,
                "section_organization": <0-25>,
                "readability": <0-25>
            }},
            "ats_recommendations": [
                "specific recommendation 1",
                "specific recommendation 2"
            ]
        }}
        
        Consider factors like:
        - Standard section headers (Experience, Education, Skills)
        - Keyword relevance and density
        - Formatting simplicity
        - Contact information completeness
        - Quantifiable achievements
        """
        
        response = self.llm._call(prompt)
        try:
            # Extract JSON from response
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                return {"ats_score": 0, "error": "Could not parse ATS analysis"}
        except:
            return {"ats_score": 0, "error": "Failed to analyze ATS score"}  
  
    def analyze_keywords(self, resume_text: str) -> Dict[str, Any]:
        """Analyze keywords and their relevance"""
        prompt = f"""
        Analyze the keywords in this resume and provide detailed keyword analysis.
        
        Resume Text:
        {resume_text}
        
        Return JSON response:
        {{
            "keyword_analysis": {{
                "technical_skills": ["skill1", "skill2"],
                "soft_skills": ["skill1", "skill2"],
                "industry_keywords": ["keyword1", "keyword2"],
                "missing_keywords": ["missing1", "missing2"],
                "keyword_density_score": <0-100>
            }},
            "keyword_recommendations": [
                "Add more industry-specific keywords",
                "Include relevant technical certifications"
            ]
        }}
        
        Focus on:
        - Technical skills and tools
        - Industry-specific terminology
        - Soft skills and competencies
        - Missing high-impact keywords
        - Keyword frequency and placement
        """
        
        response = self.llm._call(prompt)
        try:
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                return {"keyword_analysis": {}, "error": "Could not parse keyword analysis"}
        except:
            return {"keyword_analysis": {}, "error": "Failed to analyze keywords"}
    
    def analyze_action_verbs(self, resume_text: str) -> Dict[str, Any]:
        """Analyze action verbs and their strength"""
        prompt = f"""
        Analyze the action verbs used in this resume and evaluate their strength and impact.
        
        Resume Text:
        {resume_text}
        
        Return JSON response:
        {{
            "action_verb_analysis": {{
                "strong_verbs": ["achieved", "implemented", "optimized"],
                "weak_verbs": ["responsible for", "worked on", "helped"],
                "verb_strength_score": <0-100>,
                "verb_variety_score": <0-100>
            }},
            "verb_recommendations": [
                "Replace 'responsible for' with 'managed' or 'led'",
                "Use more quantifiable action verbs"
            ]
        }}
        
        Evaluate:
        - Strength and impact of action verbs
        - Variety and diversity of verbs used
        - Passive vs active voice usage
        - Industry-appropriate terminology
        """
        
        response = self.llm._call(prompt)
        try:
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                return {"action_verb_analysis": {}, "error": "Could not parse action verb analysis"}
        except:
            return {"action_verb_analysis": {}, "error": "Failed to analyze action verbs"}
    
    def analyze_strengths_weaknesses(self, resume_text: str) -> Dict[str, Any]:
        """Analyze resume strengths and weaknesses"""
        prompt = f"""
        Conduct a comprehensive analysis of this resume's strengths and weaknesses.
        
        Resume Text:
        {resume_text}
        
        Return JSON response:
        {{
            "strengths": [
                "Clear quantifiable achievements",
                "Strong technical skill set",
                "Relevant work experience"
            ],
            "weaknesses": [
                "Missing contact information",
                "Lack of quantified results",
                "Poor formatting consistency"
            ],
            "overall_impression": "Brief overall assessment",
            "competitiveness_score": <0-100>
        }}
        
        Analyze:
        - Content quality and relevance
        - Achievement quantification
        - Professional presentation
        - Completeness of information
        - Market competitiveness
        """
        
        response = self.llm._call(prompt)
        try:
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                return {"strengths": [], "weaknesses": [], "error": "Could not parse strengths/weaknesses"}
        except:
            return {"strengths": [], "weaknesses": [], "error": "Failed to analyze strengths/weaknesses"}    

    def get_improvement_recommendations(self, resume_text: str, analysis_results: Dict) -> Dict[str, Any]:
        """Generate comprehensive improvement recommendations"""
        prompt = f"""
        Based on this resume and analysis results, provide detailed improvement recommendations.
        
        Resume Text:
        {resume_text}
        
        Analysis Results:
        {json.dumps(analysis_results, indent=2)}
        
        Return JSON response:
        {{
            "priority_improvements": [
                {{
                    "category": "Content",
                    "issue": "Missing quantified achievements",
                    "recommendation": "Add specific metrics and numbers to demonstrate impact",
                    "priority": "High",
                    "example": "Increased sales by 25% over 6 months"
                }}
            ],
            "formatting_improvements": [
                "Use consistent bullet points",
                "Ensure proper section headers"
            ],
            "content_improvements": [
                "Add more industry-specific keywords",
                "Include relevant certifications"
            ],
            "overall_strategy": "Brief strategic advice for resume improvement"
        }}
        
        Focus on:
        - High-impact changes
        - ATS optimization
        - Content enhancement
        - Professional presentation
        - Market positioning
        """
        
        response = self.llm._call(prompt)
        try:
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                return {"priority_improvements": [], "error": "Could not parse recommendations"}
        except:
            return {"priority_improvements": [], "error": "Failed to generate recommendations"}
    
    def analyze_resume(self, resume_text: str) -> Dict[str, Any]:
        """Main method to analyze resume comprehensively"""
        try:
            print("🔍 Analyzing resume components...")
            
            # Perform all analyses
            ats_analysis = self.calculate_ats_score(resume_text)
            keyword_analysis = self.analyze_keywords(resume_text)
            action_verb_analysis = self.analyze_action_verbs(resume_text)
            strength_weakness_analysis = self.analyze_strengths_weaknesses(resume_text)
            
            # Combine all results
            combined_analysis = {
                "ats_analysis": ats_analysis,
                "keyword_analysis": keyword_analysis,
                "action_verb_analysis": action_verb_analysis,
                "strength_weakness_analysis": strength_weakness_analysis
            }
            
            # Get improvement recommendations
            print("💡 Generating improvement recommendations...")
            recommendations = self.get_improvement_recommendations(resume_text, combined_analysis)
            
            # Compile final results
            final_results = {
                "resume_analysis": {
                    "analysis_timestamp": "2025-01-08",
                    "ats_score": ats_analysis.get("ats_score", 0),
                    "ats_details": ats_analysis,
                    "keyword_analysis": keyword_analysis,
                    "action_verb_analysis": action_verb_analysis,
                    "strengths_weaknesses": strength_weakness_analysis,
                    "improvement_recommendations": recommendations,
                    "overall_grade": self.calculate_overall_grade(combined_analysis)
                }
            }
            
            return final_results
            
        except Exception as e:
            return {"error": f"Analysis failed: {str(e)}"}
    
    def calculate_overall_grade(self, analysis: Dict) -> str:
        """Calculate overall resume grade"""
        try:
            ats_score = analysis.get("ats_analysis", {}).get("ats_score", 0)
            keyword_score = analysis.get("keyword_analysis", {}).get("keyword_analysis", {}).get("keyword_density_score", 0)
            verb_score = analysis.get("action_verb_analysis", {}).get("action_verb_analysis", {}).get("verb_strength_score", 0)
            competitiveness = analysis.get("strength_weakness_analysis", {}).get("competitiveness_score", 0)
            
            avg_score = (ats_score + keyword_score + verb_score + competitiveness) / 4
            
            if avg_score >= 90:
                return "A+ (Excellent)"
            elif avg_score >= 80:
                return "A (Very Good)"
            elif avg_score >= 70:
                return "B (Good)"
            elif avg_score >= 60:
                return "C (Average)"
            else:
                return "D (Needs Improvement)"
        except:
            return "Unable to calculate grade"


# --- 4. Core Analysis Engine ---

class ResumeAnalyzer:
    """
    Main class that handles both job matching and company verification.
    Optimized for speed and accuracy.
    """
    
    def __init__(self):
        """Initialize the analyzer with optimized settings."""
        if not os.getenv("GROQ_API_KEY"):
            raise ValueError("GROQ_API_KEY is missing in the environment.")
            
        self.llm = ChatGroq(
            model="llama3-8b-8192",  # Fast and reliable model
            groq_api_key=os.getenv("GROQ_API_KEY"),
            temperature=0.1
        )
        
        self.search_tool = TavilySearch(tavily_api_key=os.getenv("TAVILY_API_KEY"))

    def analyze_resume(self, resume_text: str, job_description: str) -> Dict[str, Any]:
        """
        Complete resume analysis including company verification and job matching.
        """
        try:
            print("🔄 Starting company verification...")
            company_verification = self._verify_companies(resume_text)
            
            print("🔄 Starting job match analysis...")
            job_match = self._evaluate_job_match(resume_text, job_description)
            
            return {
                "company_verification": company_verification,
                "job_match_analysis": job_match
            }
            
        except Exception as e:
            return {
                "error": "Analysis failed",
                "details": str(e),
                "company_verification": [],
                "job_match_analysis": {"match_score": 0, "summary": "Error occurred"}
            }

    def _evaluate_job_match(self, resume_text: str, job_description: str) -> Dict[str, Any]:
        """Evaluate how well the resume matches the job description and provide comprehensive assessment."""
        prompt = f"""
        You are an expert HR analyst. Evaluate how well this resume matches the given job description and provide a concise assessment.

        Job Description:
        {job_description}

        Resume:
        {resume_text}

        Analyze the resume against the job requirements and provide:
        1. A numerical score from 0 to 100 (where 100 is a perfect match)
        2. A concise summary in 2-3 lines that includes:
           - Key strengths of the candidate that match the job requirements
           - Main weaknesses or gaps compared to the job requirements
           - Overall assessment of fit

        Consider:
        - Required skills and technologies alignment
        - Years of experience vs requirements
        - Educational background relevance
        - Relevant projects and achievements
        - Specific gaps or missing requirements

        Respond in exactly this format:
        Score: [your numerical score from 0-100]
        Summary: [2-3 line concise analysis covering strengths, weaknesses, and overall fit]
        """
        
        try:
            response = self.llm.invoke([HumanMessage(content=prompt)])
            output = response.content.strip()
            
            # Let the agent parse its own response
            parse_prompt = f"""
            Extract the score and summary from this analysis:

            {output}

            Return in exactly this format:
            SCORE: [number only]
            SUMMARY: [complete summary text only - ensure it's 2-3 lines covering strengths and weaknesses]
            """
            
            parse_response = self.llm.invoke([HumanMessage(content=parse_prompt)])
            parsed_output = parse_response.content.strip()
            
            # Extract score and summary from agent's parsed response
            score = 0
            summary = "Unable to evaluate"
            
            for line in parsed_output.split('\n'):
                line = line.strip()
                if line.startswith('SCORE:'):
                    try:
                        score_text = line.replace('SCORE:', '').strip()
                        score = float(score_text)
                    except (ValueError, TypeError):
                        score = 0
                elif line.startswith('SUMMARY:'):
                    summary = line.replace('SUMMARY:', '').strip()
            
            # If parsing failed, ask agent to provide components individually
            if score == 0 and summary == "Unable to evaluate":
                score_prompt = f"From this analysis, what is the numerical score (0-100)? Return only the number: {output}"
                score_response = self.llm.invoke([HumanMessage(content=score_prompt)])
                try:
                    score = float(score_response.content.strip())
                except:
                    score = 50
                
                summary_prompt = f"From this analysis, provide a concise 2-3 line summary covering candidate strengths, weaknesses, and overall fit: {output}"
                summary_response = self.llm.invoke([HumanMessage(content=summary_prompt)])
                summary = summary_response.content.strip()
            
            return {
                "match_score": score,
                "summary": summary
            }
            
        except Exception as e:
            return {
                "match_score": 0,
                "summary": f"Error evaluating job match: {str(e)}"
            }

    def _extract_companies(self, resume_text: str) -> List[str]:
        """
        Extract company names ONLY from work experience and internships using pure agent approach.
        No regex or hardcoded logic - only LLM analysis.
        """
        prompt = f"""
        You are an expert HR analyst. Analyze this resume text and extract ONLY company names where the person worked or did internships.

        STRICT RULES:
        1. Extract companies ONLY from work experience and internship sections
        2. Do NOT include educational institutions (universities, colleges, institutes, schools)
        3. Do NOT include API services, tools, or technologies (like Gemini, Meta, Brevo, Cloudinary)
        4. Do NOT include workshop providers, training centers, or certificate programs
        5. Do NOT include URLs, links, or web addresses (like "LINK", "link", "http://", etc.)
        6. Do NOT include project names, repository names, or GitHub links
        7. Look for actual employment relationships with real company entities

        Resume Text:
        {resume_text}

        Analyze the resume carefully and return only the actual company names where this person was employed or did internships. Ignore any URLs, links, or web addresses that might appear in project descriptions.

        Return each company name on a separate line with no additional text, numbers, or explanations.

        If you find companies, list them like this:
        CompanyName1
        CompanyName2
        CompanyName3

        If no companies are found, return: NO_COMPANIES_FOUND
        """
        
        try:
            response = self.llm.invoke([HumanMessage(content=prompt)])
            companies_text = response.content.strip()
            
            if not companies_text or companies_text == "NO_COMPANIES_FOUND":
                print("🏢 No companies found by agent")
                return []
            
            # Split by lines and clean
            lines = companies_text.split('\n')
            companies = []
            
            for line in lines:
                line = line.strip()
                
                # Skip empty lines
                if not line:
                    continue
                
                # Remove any numbering or bullet points that might be added
                line = line.replace('1. ', '').replace('2. ', '').replace('3. ', '')
                line = line.replace('- ', '').replace('• ', '').replace('* ', '')
                line = line.strip()
                
                # Skip if it's still empty or too short
                if len(line) < 3:
                    continue
                
                # Skip obvious agent artifacts and URLs/links
                if any(phrase in line.lower() for phrase in [
                    'no companies', 'found', 'based on', 'analysis', 'resume',
                    'link', 'http', 'www', '.com', '.org', '.net', 'github',
                    'url', 'website', 'repository', 'repo'
                ]):
                    continue
                
                companies.append(line)
            
            # Remove duplicates while preserving order
            unique_companies = []
            for company in companies:
                if company not in unique_companies:
                    unique_companies.append(company)
            
            print(f"🏢 Extracted companies (agent): {unique_companies}")
            return unique_companies[:3]  # Limit to 3 companies for performance
            
        except Exception as e:
            print(f"❌ Error extracting companies: {e}")
            return []

    def _verify_companies(self, resume_text: str) -> List[Dict[str, Any]]:
        """Verify the legitimacy of extracted companies using pure agent analysis."""
        companies = self._extract_companies(resume_text)
        verification_results = []
        
        for company in companies:
            try:
                print(f"🔍 Verifying: {company}")
                
                # Get search results for company information
                try:
                    search_results = self.search_tool.invoke(f"{company} company reviews legitimacy employee payment")
                    search_text = str(search_results) if search_results else "No search results available"
                except Exception as search_error:
                    print(f"⚠️ Search failed for {company}: {search_error}")
                    search_text = "Search unavailable"
                
                # Let the agent analyze the company comprehensively
                verification_prompt = f"""
                You are an expert company analyst. Analyze the company '{company}' based on the following information and provide a comprehensive verification report.

                Available Information:
                {search_text}

                Analyze and determine:
                1. Legitimacy Status: Is this a legitimate company? (Legitimate/Suspicious/Unknown)
                2. Payment Status: How does this company handle employee/intern payments? (Paid/Unpaid/Student Paid Fee/Unknown)
                3. Key Findings: What are the most important findings about this company?
                4. Reputation Score: Rate the company's reputation from 1-10 (1=very poor, 10=excellent)

                Provide your analysis in exactly this format:
                STATUS: [Legitimate/Suspicious/Unknown]
                PAYMENT: [Paid/Unpaid/Student Paid Fee/Unknown]
                FINDINGS: [One sentence summary of key findings]
                SCORE: [number from 1-10]
                """
                
                response = self.llm.invoke([HumanMessage(content=verification_prompt)])
                result_text = response.content.strip()
                
                # Let the agent parse its own response
                parse_prompt = f"""
                Extract the specific information from this verification report:

                {result_text}

                Return in exactly this format with only the requested values:
                STATUS: [status value only]
                PAYMENT: [payment value only]  
                FINDINGS: [findings text only]
                SCORE: [score number only]
                """
                
                parse_response = self.llm.invoke([HumanMessage(content=parse_prompt)])
                parsed_result = parse_response.content.strip()
                
                # Extract values from agent's parsed response
                status = "Unknown"
                payment = "Unknown"
                summary = "Limited information available"
                score = 5
                
                for line in parsed_result.split('\n'):
                    line = line.strip()
                    if line.startswith('STATUS:'):
                        status = line.replace('STATUS:', '').strip()
                    elif line.startswith('PAYMENT:'):
                        payment = line.replace('PAYMENT:', '').strip()
                    elif line.startswith('FINDINGS:'):
                        summary = line.replace('FINDINGS:', '').strip()
                    elif line.startswith('SCORE:'):
                        try:
                            score = int(line.replace('SCORE:', '').strip())
                        except (ValueError, TypeError):
                            score = 5
                
                # If parsing failed, ask agent for individual components
                if status == "Unknown" and payment == "Unknown" and summary == "Limited information available":
                    status_prompt = f"From this analysis of {company}, what is the legitimacy status? Answer only: Legitimate, Suspicious, or Unknown: {result_text}"
                    status_response = self.llm.invoke([HumanMessage(content=status_prompt)])
                    status = status_response.content.strip()
                    
                    payment_prompt = f"From this analysis of {company}, what is the payment status? Answer only: Paid, Unpaid, Student Paid Fee, or Unknown: {result_text}"
                    payment_response = self.llm.invoke([HumanMessage(content=payment_prompt)])
                    payment = payment_response.content.strip()
                    
                    findings_prompt = f"From this analysis of {company}, provide a one-sentence summary of key findings: {result_text}"
                    findings_response = self.llm.invoke([HumanMessage(content=findings_prompt)])
                    summary = findings_response.content.strip()
                    
                    score_prompt = f"From this analysis of {company}, what reputation score (1-10) would you give? Answer only with the number: {result_text}"
                    score_response = self.llm.invoke([HumanMessage(content=score_prompt)])
                    try:
                        score = int(score_response.content.strip())
                    except:
                        score = 5
                
                verification_results.append({
                    "company_name": company,
                    "legitimacy_status": status,
                    "payment_status": payment,
                    "key_findings": summary,
                    "reputation_score": score
                })
                
                print(f"✅ Verified: {company} - {status}")
                
            except Exception as e:
                print(f"❌ Error verifying {company}: {e}")
                verification_results.append({
                    "company_name": company,
                    "legitimacy_status": "Error",
                    "payment_status": "Unknown",
                    "key_findings": f"Verification failed: {str(e)}",
                    "reputation_score": 0
                })
        
        return verification_results

# --- 5. FastAPI Endpoints ---

@app.get("/")
async def root():
    """Root endpoint with API information."""
    return {
        "message": "Resume Analyzer API",
        "version": "1.0.0",
        "endpoints": {
            "/analyze": "POST - Upload resume and job description for analysis",
            "/analyze-text": "POST - Analyze text directly",
            "/detailed-resume-analysis": "POST - Detailed resume analysis with ATS scoring",
            "/company-legitimacy-check": "POST - Check company legitimacy",
            "/health": "GET - Health check",
            "/docs": "GET - API documentation"
        }
    }

@app.get("/health")
async def health_check():
    """Health check endpoint."""
    try:
        groq_key = os.getenv("GROQ_API_KEY")
        tavily_key = os.getenv("TAVILY_API_KEY")
        
        if not groq_key:
            return {"status": "unhealthy", "message": "GROQ_API_KEY not configured"}
        if not tavily_key:
            return {"status": "unhealthy", "message": "TAVILY_API_KEY not configured"}
        
        return {"status": "healthy", "message": "All systems operational"}
    except Exception as e:
        return {"status": "unhealthy", "message": f"Health check failed: {str(e)}"}

@app.post("/analyze", response_model=AnalysisResponse)
async def analyze_files(
    resume_file: UploadFile = File(...),
    job_description_file: UploadFile = File(...)
):
    """Analyze resume and job description files."""
    try:
        # Validate file types
        allowed_extensions = ['.pdf', '.docx']
        
        if not any(resume_file.filename.lower().endswith(ext) for ext in allowed_extensions):
            raise HTTPException(status_code=400, detail=f"Resume file must be PDF or DOCX")
        
        if not any(job_description_file.filename.lower().endswith(ext) for ext in allowed_extensions):
            raise HTTPException(status_code=400, detail=f"Job description file must be PDF or DOCX")
        
        # Read file contents
        resume_content = await resume_file.read()
        job_description_content = await job_description_file.read()
        
        # Extract text
        resume_text = extract_text_from_file_bytes(resume_content, resume_file.filename)
        job_description_text = extract_text_from_file_bytes(job_description_content, job_description_file.filename)
        
        if resume_text.startswith("Error"):
            raise HTTPException(status_code=400, detail=f"Resume processing failed: {resume_text}")
        
        if job_description_text.startswith("Error"):
            raise HTTPException(status_code=400, detail=f"Job description processing failed: {job_description_text}")
        
        # Run analysis
        analyzer = ResumeAnalyzer()
        result = analyzer.analyze_resume(resume_text, job_description_text)
        
        if "error" in result:
            raise HTTPException(status_code=500, detail=result["details"])
        
        return AnalysisResponse(**result)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

@app.post("/analyze-text", response_model=AnalysisResponse)
async def analyze_text(request: TextAnalysisRequest):
    """Analyze resume and job description text."""
    try:
        if len(request.resume_text.strip()) < 50:
            raise HTTPException(status_code=400, detail="Resume text too short")
        
        if len(request.job_description_text.strip()) < 20:
            raise HTTPException(status_code=400, detail="Job description text too short")
        
        # Run analysis
        analyzer = ResumeAnalyzer()
        result = analyzer.analyze_resume(request.resume_text, request.job_description_text)
        
        if "error" in result:
            raise HTTPException(status_code=500, detail=result["details"])
        
        return AnalysisResponse(**result)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

@app.post("/detailed-resume-analysis")
async def detailed_resume_analysis(request: ResumeAnalysisRequest):
    """Perform detailed resume analysis with ATS scoring."""
    try:
        if len(request.resume_text.strip()) < 50:
            raise HTTPException(status_code=400, detail="Resume text too short")
        
        # Get API key from environment
        api_key = os.getenv("MOONSHOT_API_KEY") or os.getenv("GROQ_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="API key not configured")
        
        # Initialize detailed analyzer
        analyzer = DetailedResumeAnalyzer(api_key=api_key)
        
        # Analyze resume
        result = analyzer.analyze_resume(request.resume_text)
        
        if "error" in result:
            raise HTTPException(status_code=500, detail=result["error"])
        
        return JSONResponse(content=result)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Detailed analysis failed: {str(e)}")

@app.post("/detailed-resume-analysis-file")
async def detailed_resume_analysis_file(resume_file: UploadFile = File(...)):
    """Perform detailed resume analysis on uploaded file."""
    try:
        # Validate file type
        allowed_extensions = ['.pdf', '.docx']
        
        if not any(resume_file.filename.lower().endswith(ext) for ext in allowed_extensions):
            raise HTTPException(status_code=400, detail=f"Resume file must be PDF or DOCX")
        
        # Read file content
        resume_content = await resume_file.read()
        
        # Extract text
        resume_text = extract_text_from_file_bytes(resume_content, resume_file.filename)
        
        if resume_text.startswith("Error"):
            raise HTTPException(status_code=400, detail=f"Resume processing failed: {resume_text}")
        
        # Get API key from environment
        api_key = os.getenv("MOONSHOT_API_KEY") or os.getenv("GROQ_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="API key not configured")
        
        # Initialize detailed analyzer
        analyzer = DetailedResumeAnalyzer(api_key=api_key)
        
        # Analyze resume
        result = analyzer.analyze_resume(resume_text)
        
        if "error" in result:
            raise HTTPException(status_code=500, detail=result["error"])
        
        return JSONResponse(content=result)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Detailed analysis failed: {str(e)}")

@app.post("/company-legitimacy-check", response_model=CompanyLegitimacyResult)
async def company_legitimacy_check(request: CompanyLegitimacyRequest):
    """Check company legitimacy."""
    try:
        if len(request.company_name.strip()) < 2:
            raise HTTPException(status_code=400, detail="Company name too short")
        
        # Initialize company checker
        checker = CompanyLegitimacyChecker()
        
        # Check company legitimacy
        result = checker.check_company_legitimacy(request.company_name, request.job_role)
        
        if result["status"] == "error":
            raise HTTPException(status_code=500, detail=result["assessment"])
        
        return CompanyLegitimacyResult(**result)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Company legitimacy check failed: {str(e)}")

# --- 6. Main Application Functions ---

def run_complete_analysis(resume_path: str, job_description_path: str) -> Dict[str, Any]:
    """Run complete analysis on resume and job description files."""
    print(f"📄 Processing resume file: {resume_path}")
    print(f"📄 Processing job description file: {job_description_path}")
    
    # Extract text from resume file
    resume_text = process_resume_file(resume_path)
    
    if resume_text.startswith("Error"):
        return {"error": "Resume file processing failed", "details": resume_text}
    
    print(f"✅ Extracted {len(resume_text)} characters from resume file")
    
    # Extract text from job description file
    job_description_text = process_resume_file(job_description_path)
    
    if job_description_text.startswith("Error"):
        return {"error": "Job description file processing failed", "details": job_description_text}
    
    print(f"✅ Extracted {len(job_description_text)} characters from job description file")
    
    # Run analysis
    analyzer = ResumeAnalyzer()
    result = analyzer.analyze_resume(resume_text, job_description_text)
    
    return result

def run_text_analysis(resume_text: str, job_description_text: str) -> Dict[str, Any]:
    """Run analysis on resume and job description text directly."""
    print("📝 Processing resume and job description text...")
    
    analyzer = ResumeAnalyzer()
    result = analyzer.analyze_resume(resume_text, job_description_text)
    
    return result

def run_detailed_resume_analysis(resume_path: str) -> Dict[str, Any]:
    """Run detailed resume analysis with ATS scoring."""
    print(f"📄 Processing resume file for detailed analysis: {resume_path}")
    
    # Get API key from environment
    api_key = os.getenv("MOONSHOT_API_KEY") or os.getenv("GROQ_API_KEY")
    if not api_key:
        return {"error": "API key not configured"}
    
    try:
        # Initialize analyzer
        analyzer = DetailedResumeAnalyzer(api_key=api_key)
        
        # Extract text
        resume_text = process_resume_file(resume_path)
        
        if resume_text.startswith("Error"):
            return {"error": "Resume file processing failed", "details": resume_text}
        
        print(f"✅ Extracted {len(resume_text)} characters from resume file")
        
        # Analyze resume
        result = analyzer.analyze_resume(resume_text)
        
        return result
        
    except Exception as e:
        return {"error": f"Detailed analysis failed: {str(e)}"}

def run_company_legitimacy_check(company_name: str, job_role: str = None) -> Dict[str, Any]:
    """Run company legitimacy check."""
    print(f"🔍 Checking company legitimacy: {company_name}")
    if job_role:
        print(f"📋 Job role: {job_role}")
    
    try:
        # Initialize checker
        checker = CompanyLegitimacyChecker()
        
        # Check company legitimacy
        result = checker.check_company_legitimacy(company_name, job_role)
        
        return result
        
    except Exception as e:
        return {"error": f"Company legitimacy check failed: {str(e)}"}

def save_analysis_results(results: Dict[str, Any], filename: str) -> None:
    """Save analysis results to JSON file."""
    try:
        with open(filename, 'w') as f:
            json.dump(results, f, indent=4)
        print(f"💾 Results saved to: {filename}")
    except Exception as e:
        print(f"❌ Error saving results: {str(e)}")

# --- 7. Main Execution Block ---

if __name__ == "__main__":
    import sys
    
    # Check command line arguments for different modes
    if len(sys.argv) > 1:
        mode = sys.argv[1]
        
        if mode == "--api":
            # Run FastAPI server
            uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)
            
        elif mode == "--detailed-analysis" and len(sys.argv) > 2:
            # Run detailed resume analysis
            resume_path = sys.argv[2]
            print("🚀 Starting Detailed Resume Analysis...")
            print("-" * 60)
            
            try:
                result = run_detailed_resume_analysis(resume_path)
                
                print("✅ Detailed Analysis Complete!")
                print("-" * 60)
                print("📊 Analysis Results:")
                print(json.dumps(result, indent=4))
                
                # Save to file
                output_file = "detailed_resume_analysis.json"
                save_analysis_results(result, output_file)
                
            except Exception as e:
                print(f"❌ Error: {str(e)}")
                
        elif mode == "--company-check" and len(sys.argv) > 2:
            # Run company legitimacy check
            company_name = sys.argv[2]
            job_role = sys.argv[3] if len(sys.argv) > 3 else None
            
            print("🚀 Starting Company Legitimacy Check...")
            print("-" * 60)
            
            try:
                result = run_company_legitimacy_check(company_name, job_role)
                
                print("✅ Company Check Complete!")
                print("-" * 60)
                print("📊 Check Results:")
                print(json.dumps(result, indent=4))
                
                # Save to file
                filename_parts = [company_name.replace(' ', '_').replace('/', '_')]
                if job_role:
                    filename_parts.append(job_role.replace(' ', '_').replace('/', '_'))
                output_file = f"{'_'.join(filename_parts)}_legitimacy_check.json"
                save_analysis_results(result, output_file)
                
            except Exception as e:
                print(f"❌ Error: {str(e)}")
                
        elif mode == "--help":
            print("🚀 Resume Analyzer - Usage Guide")
            print("=" * 50)
            print("Available modes:")
            print("  --api                           : Start FastAPI server")
            print("  --detailed-analysis <file>      : Run detailed resume analysis")
            print("  --company-check <name> [role]   : Check company legitimacy")
            print("  --help                          : Show this help message")
            print()
            print("Default mode (no arguments): Run basic resume + job description analysis")
            print()
            print("Examples:")
            print("  python app.py --api")
            print("  python app.py --detailed-analysis /path/to/resume.pdf")
            print("  python app.py --company-check 'Google' 'Software Engineer'")
            
        else:
            print("❌ Invalid arguments. Use --help for usage guide.")
            
    else:
        # Default mode: Run original script logic
        RESUME_FILE_PATH = "/Users/apple/Desktop/employers/Resume.pdf"
        JOB_DESCRIPTION_FILEPATH = "/Users/apple/Desktop/employers/Untitled document (1).pdf"
        
        print("🚀 Initializing the Optimized Resume Analyzer...")
        print("-" * 60)
        
        try:
            print(f"📁 Processing files:")
            print(f"   Resume: {RESUME_FILE_PATH}")
            print(f"   Job Description: {JOB_DESCRIPTION_FILEPATH}")
            
            final_report = run_complete_analysis(RESUME_FILE_PATH, JOB_DESCRIPTION_FILEPATH)
            
            print("✅ Analysis Complete!")
            print("-" * 60)
            print("📊 Final Report:")
            print(json.dumps(final_report, indent=4))
            
            # Save to file
            output_file = "complete_analysis_report.json"
            save_analysis_results(final_report, output_file)
            
        except Exception as e:
            print(f"❌ Error: {str(e)}")
            print("Please check your API keys and file paths.")
